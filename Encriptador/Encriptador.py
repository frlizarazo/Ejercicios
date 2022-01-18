# -*- coding: utf-8 -*-
"""
Created on Sun Jan 16 23:30:16 2022

@author: Frank
"""

# import docx as d
# import numpy as n
import random as r
import os
import tkinter.filedialog as f
import tkinter as t
import docx2txt as d
from docx import Document

script=''
try:
    os.makedirs('Files/')
except FileExistsError:
    pass

def read():
    global script
    file   = str(f.askopenfile(parent=root,title='Seleccionar Texto',
                           filetypes=(('Archivo Word','*.docx'),),
                           initialdir=(os.getcwd())))
    file   = file[file.find("'")+1:file.find("'",file.find("'")+1)]
    
    script = d.process(file)

def encrypt():
    global newScript, key
    
    key = []
    newScript= ''
   

    for n in script:
        gap        = r.randint(65,90)
        char       = chr(ord(n)+gap)
        newScript += char
        key.append(chr(gap))
    
    key=''.join(key)
    document=Document()
    document.add_paragraph(newScript)
    document.save('Files\Encriptado.docx')   
    
    docKey=Document()
    docKey.add_paragraph(key)
    docKey.save('Files\key.docx')
    
def decrypt():
    global oldScript
    file   = str(f.askopenfile(parent=root,title='Seleccionar Clave',
                           filetypes=(('Archivo Word','*.docx'),),
                           initialdir=(os.getcwd())))
    file   = file[file.find("'")+1:file.find("'",file.find("'")+1)]
    
    key = d.process(file)
    
    i=0;
    oldScript=''
    for n in newScript:
        char       = chr(ord(n)-ord(key[i]))
        oldScript += char
        i         +=1
        
    document=Document()
    document.add_paragraph(oldScript)
    document.save('Files\Desencriptado.docx')
    
def close():
    root.destroy()

root=t.Tk()
root.title('Codificar')

button1 = t.Button(root,text='Examinar',command=read).grid(row=0,column=0,padx=[20,0],sticky='ew')
button2 = t.Button(root,text='Encriptar',command=encrypt).grid(row=1,column=0,padx=[20,0],sticky='ew')
button3 = t.Button(root,text='Desencriptar',command=decrypt).grid(row=2,column=0,padx=[20,0],sticky='ew')
button4 = t.Button(root,text='Cerrar',command=close).grid(row=3,column=0,padx=[20,0],sticky='ew')

root.mainloop()